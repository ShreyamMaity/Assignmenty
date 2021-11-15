from __future__ import print_function
import os
import base64
import os.path
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import subprocess
import mimetypes
import docx
import glob
from datetime import datetime
import re

def auth():
    SCOPES =  ['https://www.googleapis.com/auth/gmail.send']
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.json', 'w') as token:
            token.write(creds.to_json())

    service = build('gmail', 'v1', credentials=creds)
    return service

def mailSender(assignment, subject, body, mail):
    service = auth()
    file_attachements = [r'./'+ assignment +'.docx']
    emailMsg = subject
    Message = MIMEMultipart()
    Message['to'] = mail
    Message['subject'] = body
    Message.attach(MIMEText(emailMsg, 'plain'))

    for attachement in file_attachements:
        content_type, encoding = mimetypes.guess_type(attachement)
        main_type, sub_type = content_type.split('/', 1)
        file_name = os.path.basename(attachement)

        f = open(attachement, 'rb')

        myFile = MIMEBase(main_type, sub_type)
        myFile.set_payload(f.read())
        myFile.add_header('Content-Disposition',
                              'attachement', filename=file_name)
        encoders.encode_base64(myFile)

        f.close()

        Message.attach(myFile)
    raw_string = base64.urlsafe_b64encode(Message.as_bytes()).decode()
    message = service.users().messages().send(userId='me', body={'raw': raw_string}).execute()
    print('Message Id: %s' % message['id'])

def getCommitFiles():
    files = subprocess.check_output(['git', 'diff-tree', '--no-commit-id', '--name-only', '-r', 'HEAD']).decode('utf-8').split('\n')
    return files

def getPyFiles(fileList):
    pyFiles = []
    sort = {}
    for file in fileList:
        if file.endswith('.py'):
            m = re.search(r'\d+$', file.replace('.py', ''))
            if m:
                sort[m.group()] = file
    for k, v in sorted(sort.items(), key=lambda x: int(x[0])):
        pyFiles.append(v)
    return pyFiles

def getQuestions(fileList):
    questions = {}
    for file in fileList:
        with open(file, 'r') as f:
            for line in f:
                if line.startswith('\'\'\''):
                    questions[re.search(r'\d+$', file.replace('.py', '')).group()] = line.replace('\'\'\'', '').strip()
    return questions

def getRunSnaps(fileList):
    runSnaps = []
    sort = {}
    try : 
        for file in fileList:
            if file.endswith('.png'):
                m = re.search(r'\d+$', file.replace('.png', ''))
                if m:
                    sort[m.group()] = file
        for k, v in sorted(sort.items(), key=lambda x: int(x[0])):
            runSnaps.append(v)
        return runSnaps
    except:
        for file in fileList:
            if file.endswith('.jpg'):
                m = re.search(r'\d+$', file.replace('.jpg', ''))
                if m:
                    sort[m.group()] = file
        for k, v in sorted(sort.items(), key=lambda x: int(x[0])):
            runSnaps.append(v)
        return runSnaps

def getTime():
    return datetime.now().strftime("%d/%m/%Y")

def buildPdf(assignment, questions, snaps, template, fileName):
    doc = docx.Document()
    base = docx.Document(template)
    for para in base.paragraphs:
        if para.text.find('Date of')!=-1:
            para.text += getTime()
        bc = doc.add_paragraph(para.text)
        bc.bold = True
        bc.alignment = para.alignment
        bc.style = para.style  
    doc.add_paragraph("_________________________________________________________________________________________________________")
    for key, value in questions.items():
        quesNo = doc.add_paragraph()
        quesNo.add_run("Question No: ").bold = True
        quesNo.add_run(key)
        quesNo.style = 'List Number'
        
        ques = doc.add_paragraph()
        ques.add_run("Question: \t").bold = True
        ques.add_run(value)
        ques.style = 'List Bullet'
        for code in assignment:
            if key in code:
                with open(code, 'r') as f:
                    data = f.read()
                code = doc.add_paragraph()
                code.add_run("Code: \n\n").bold = True
                code.add_run(data)
                code.style = 'List Bullet'
        for snap in snaps:
            if key in snap:
                tit = doc.add_paragraph()
                tit.add_run("Snapshot: \n\n").bold = True
                tit.style = 'List Bullet'
                doc.add_picture(snap, width=docx.shared.Cm(16),height=docx.shared.Inches(2))
                
        doc.add_paragraph("_________________________________________________________________________________________________________")
    doc.add_paragraph('\n\nAuto Generated with Assignmenty Bot. Check me out on github: github.com/ShreyamMaity/Assignmenty').alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.save(fileName + '.docx')

def removeFiles(pdfName):
    if os.path.exists(pdfName + '.docx'):
        os.remove(pdfName + '.docx')
    for file in glob.glob("*.png"):
        os.remove(file)
    
def main():

    mailId = os.environ['INPUT_MAILID']
    fileName = os.environ['INPUT_FILENAME']
    fileList = getCommitFiles()
    pyFiles = getPyFiles(fileList)
    questions = getQuestions(pyFiles)
    snaps = getRunSnaps(fileList)
    buildPdf(pyFiles, questions, snaps, 'template.docx' , fileName)
    mailSender(fileName, 'Your Assignment Is here', 'Hello User, Thanks for using Assignment Bot. Hope You love our sevice 😊\n', mailId)
    removeFiles(fileName)


    

if __name__ == '__main__':
    main()
